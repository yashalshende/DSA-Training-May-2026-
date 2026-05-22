from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(r"C:\Users\shend\OneDrive\Desktop\DSA\Python_DSA_Code_Explanations.docx")


def set_cell_shading(paragraph, fill="F3F5F7"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Python DSA Code Explanations")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(27, 54, 93)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run(
        "Easy study notes for string, list, linked list, stack, queue, tree, and graph"
    )
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    intro = doc.add_paragraph()
    intro.add_run("How to use this file: ").bold = True
    intro.add_run(
        "Each section shows the original code snippet, then explains the important lines in simple language. "
        "A few broken lines are kept as they appear in the files and marked with a short note."
    )

    chapters = get_chapters()

    for chapter_index, chapter in enumerate(chapters):
        if chapter_index > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_chapter(doc, chapter)

    doc.save(OUTPUT_PATH)


def add_chapter(doc, chapter):
    heading = doc.add_paragraph()
    run = heading.add_run(chapter["title"])
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(27, 54, 93)

    if chapter.get("summary"):
        para = doc.add_paragraph()
        para.add_run("Chapter idea: ").bold = True
        para.add_run(chapter["summary"])

    for section in chapter["sections"]:
        add_section(doc, section)


def add_section(doc, section):
    doc.add_paragraph()
    head = doc.add_paragraph()
    run = head.add_run(section["title"])
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13)

    source = doc.add_paragraph()
    source.add_run("Source: ").bold = True
    source.add_run(section["source"])

    purpose = doc.add_paragraph()
    purpose.add_run("Purpose: ").bold = True
    purpose.add_run(section["purpose"])

    code_label = doc.add_paragraph()
    code_label.add_run("Code").bold = True

    code_para = doc.add_paragraph()
    set_cell_shading(code_para)
    code_run = code_para.add_run(section["code"].rstrip())
    code_run.font.name = "Consolas"
    code_run.font.size = Pt(9.5)

    explain_label = doc.add_paragraph()
    explain_label.add_run("Easy explanation").bold = True

    for item in section["explanations"]:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(2)
        para.add_run(item)

    if section.get("dry_run"):
        label = doc.add_paragraph()
        label.add_run("Short dry run").bold = True
        dry = doc.add_paragraph()
        dry.add_run(section["dry_run"])

    if section.get("note"):
        note = doc.add_paragraph()
        note.add_run("Note: ").bold = True
        note.add_run(section["note"])

    if section.get("output"):
        out = doc.add_paragraph()
        out.add_run("Output / key logic: ").bold = True
        out.add_run(section["output"])


def get_chapters():
    return [
        {
            "title": "1. String",
            "summary": "This chapter explains the string practice code from slicing.py in a short and easy way.",
            "sections": [
                {
                    "title": "String indexing and slicing",
                    "source": "slicing.py lines 1-14",
                    "purpose": "This code shows how to read characters and parts of a string using indexes and slices.",
                    "code": """# name = "prashantjha"
# print(name[0])
# print(name[1])
# print(name[-1])
# print(name[0:5])
# print(name[1:])
# print(name[:5])
# print(name[:])
# print(name[1:8:2])
# print(name[::-1])""",
                    "explanations": [
                        'Line 1 stores the word "prashantjha" in the variable name.',
                        "Line 2 prints the character at index 0, so it gives the first letter.",
                        "Line 3 prints the character at index 1, so it gives the second letter.",
                        "Line 4 uses -1, which means the last character of the string.",
                        "Line 5 prints characters from index 0 up to 4.",
                        "Line 6 prints from index 1 till the end.",
                        "Line 7 prints from the start up to index 4.",
                        "Line 8 prints the full string.",
                        "Line 9 takes a step of 2, so it skips one character each time.",
                        "Line 10 uses a negative step, so it reverses the string.",
                    ],
                    "output": "Main idea: positive indexes move left to right, negative indexes move from the end, and slicing lets you take part of a string.",
                },
                {
                    "title": "Changing string letter case",
                    "source": "slicing.py lines 16-22",
                    "purpose": "This code shows common built-in methods that change letter case.",
                    "code": """# s = "Python are High level programming Language"
# print(s.lower())
# print(s.upper())
# print(s.swapcase())
# print(s.title())
# print(s.capitalize())""",
                    "explanations": [
                        "Line 1 stores a sentence in s.",
                        "Line 2 changes all letters to small letters.",
                        "Line 3 changes all letters to capital letters.",
                        "Line 4 flips small letters to capital and capital to small.",
                        "Line 5 makes the first letter of each word capital.",
                        "Line 6 makes only the first letter of the full sentence capital.",
                    ],
                    "output": "These methods do not change the original string in place; they return a new string.",
                },
                {
                    "title": "Loop through each character",
                    "source": "slicing.py lines 33-35",
                    "purpose": "This code prints one character at a time from a string.",
                    "code": """# name = "yashal"
# for i in name:
#     print(i)""",
                    "explanations": [
                        'Line 1 stores the string "yashal".',
                        "Line 2 starts a loop and takes one character at a time from the string.",
                        "Line 3 prints the current character.",
                    ],
                    "output": "The letters are printed one below another: y, a, s, h, a, l.",
                },
                {
                    "title": "Reverse a string using a loop",
                    "source": "slicing.py lines 41-46",
                    "purpose": "This code builds a reversed string manually.",
                    "code": """# name = "yashal"
# newname = ""
# N = len(name)
# for i in range(N-1,-1,-1):
#     newname += name[i]
# print(newname)""",
                    "explanations": [
                        'Line 1 stores the word "yashal".',
                        "Line 2 starts with an empty string where the reverse will be built.",
                        "Line 3 finds the total number of characters.",
                        "Line 4 starts from the last index and moves backward to index 0.",
                        "Line 5 adds each character to newname one by one.",
                        "Line 6 prints the final reversed string.",
                    ],
                    "dry_run": 'For "yashal", the loop reads l, a, h, s, a, y and makes "lahsay".',
                    "output": 'Final output: "lahsay".',
                },
                {
                    "title": "Palindrome check",
                    "source": "slicing.py lines 48-54",
                    "purpose": "This code checks whether a string reads the same from both sides.",
                    "code": """# name = "yashal"
# print(name)
# print(name[::-1])
# if name == name[::-1]:
#     print("Palindrome")
# else:
#     print("Not Palindrome")""",
                    "explanations": [
                        "Line 1 stores the word to check.",
                        "Line 2 prints the original word.",
                        "Line 3 prints the reversed word using slicing.",
                        "Line 4 compares the original and reversed strings.",
                        "Line 5 runs if both are same.",
                        "Line 7 runs if they are different.",
                    ],
                    "output": 'If the word and its reverse match, it is a palindrome; otherwise it is not.',
                },
                {
                    "title": "Count vowels and consonants",
                    "source": "slicing.py lines 56-67",
                    "purpose": "This code counts how many vowels and consonants are present in a string.",
                    "code": """# name = "yashal"
# vowels = 0
# consonants = 0
# for i in name:
#     if i in "aeiouAEIOU":
#         vowels += 1
#     else:
#         consonants += 1
# print("Vowels:", vowels)
# print("Consonants:", consonants)""",
                    "explanations": [
                        "Line 1 stores the input word.",
                        "Lines 2 and 3 start both counters from 0.",
                        "Line 4 checks each character one by one.",
                        "Line 5 checks whether the character is a vowel.",
                        "Line 6 increases the vowel counter.",
                        "Lines 7 and 8 increase the consonant counter when the character is not a vowel.",
                        "Lines 9 and 10 print the final counts.",
                    ],
                    "output": "The program separates vowels from the remaining letters and counts both.",
                },
                {
                    "title": "Anagram check",
                    "source": "slicing.py lines 69-76",
                    "purpose": "This code checks whether two strings contain the same letters in a different order.",
                    "code": """# name1 = "yashal"
# name2 = "lahsay"
# print(sorted(name1))
# print(sorted(name2))
# if sorted(name1) == sorted(name2):
#     print("Anagram")
# else:
#     print("Not Anagram")""",
                    "explanations": [
                        "Lines 1 and 2 store the two strings.",
                        "Lines 3 and 4 sort both strings into letter order.",
                        "Line 5 compares the sorted versions.",
                        "Line 6 prints Anagram if the letters match.",
                        "Line 8 prints Not Anagram if they do not match.",
                    ],
                    "output": "If sorted letters are equal, both words are anagrams.",
                },
                {
                    "title": "Pangram check",
                    "source": "slicing.py lines 78-85",
                    "purpose": "This code checks whether a sentence contains every alphabet letter at least once.",
                    "code": """# sentence = "The quick brown fox jumps over the lazy dog"
# alphabet = "abcdefghijklmnopqrstuvwxyz"
# for char in alphabet:
#     if char not in sentence.lower():
#         print("Not Pangram")
#         break
# else:
#     print("Pangram")""",
                    "explanations": [
                        "Line 1 stores the sentence to test.",
                        "Line 2 stores all lowercase alphabet letters.",
                        "Line 3 checks each alphabet letter one by one.",
                        "Line 4 converts the sentence to lowercase and checks whether the letter exists.",
                        "Lines 5 and 6 stop early if any letter is missing.",
                        "Lines 7 and 8 run only if the loop finishes without breaking, which means every letter was found.",
                    ],
                    "output": 'The sample sentence is a pangram, so the result is "Pangram".',
                },
                {
                    "title": "Count words and reverse word order",
                    "source": "slicing.py lines 87-96",
                    "purpose": "This code uses split and join to work with words inside a sentence.",
                    "code": """# sentence = "The quick brown fox jumps over the lazy dog"
# words = sentence.split()
# print("Number of words:", len(words))
#
# sentence = "The quick brown fox jumps over the lazy dog"
# words = sentence.split()
# reversed_sentence = " ".join(reversed(words))
# print("Reversed sentence:", reversed_sentence)""",
                    "explanations": [
                        "Line 1 stores a full sentence.",
                        "Line 2 uses split() to break the sentence into a list of words.",
                        "Line 3 counts how many words are in that list.",
                        "Line 5 stores the sentence again for the next example.",
                        "Line 6 again converts the sentence into words.",
                        "Line 7 reverses the word order and joins them back with spaces.",
                        "Line 8 prints the new sentence.",
                    ],
                    "output": "split() turns a sentence into a list; join() turns a list back into a string.",
                },
                {
                    "title": "Title case and subsequence check",
                    "source": "slicing.py lines 120-131",
                    "purpose": "This code changes sentence style and checks whether one string appears in order inside another.",
                    "code": """# sentence = "the quick brown fox jumps over the lazy dog"
# title_case_sentence = sentence.title()
# print("Title Case:", title_case_sentence)
#
# def is_subsequence(s1, s2):
#     it = iter(s2)
#     return all(char in it for char in s1)
# s1 = "abc"
# s2 = "ahbgdc"
# print(is_subsequence(s1, s2))""",
                    "explanations": [
                        "Lines 1 to 3 convert a sentence to title case and print it.",
                        "Line 5 starts a function that checks subsequence logic.",
                        "Line 6 creates an iterator over s2 so characters are checked in order.",
                        "Line 7 checks whether every character of s1 appears in that order inside s2.",
                        "Lines 8 and 9 store sample values.",
                        "Line 10 prints the answer from the function.",
                    ],
                    "output": 'For "abc" and "ahbgdc", the answer is True because a, b, and c appear in order.',
                },
                {
                    "title": "Useful string checking methods",
                    "source": "slicing.py lines 134-147",
                    "purpose": "This code shows handy built-in methods used with strings.",
                    "code": """# print('prashantjha777'.isalnum())
# print('prashantjha'.isalpha())
# print('777f'.isdigit())
# print('sdsdsdsd'.islower())
# print('PRASHANTj'.isupper())
# print('My Name Is Prashant'.istitle())
# print(''.isspace())
# print("Hello".startswith("He"))
# print("Hello".endswith("lo"))
# print("yashal".find("ha"))
# print("yashal".count("a"))""",
                    "explanations": [
                        "isalnum() checks letters and digits only.",
                        "isalpha() checks letters only.",
                        "isdigit() checks digits only.",
                        "islower() checks whether all letters are lowercase.",
                        "isupper() checks whether all letters are uppercase.",
                        "istitle() checks title case format.",
                        "isspace() checks whether the string has only spaces.",
                        "startswith() checks the starting text.",
                        "endswith() checks the ending text.",
                        "find() gives the starting index of a substring.",
                        "count() tells how many times a character appears.",
                    ],
                    "output": "These methods are useful for validation, searching, and text formatting.",
                },
            ],
        },
        {
            "title": "2. List",
            "summary": "This chapter explains list behavior, especially how lists can be changed inside functions.",
            "sections": [
                {
                    "title": "List changed inside a function",
                    "source": "list.py lines 1-7",
                    "purpose": "This example shows that a list can change when we pass it to a function.",
                    "code": """# def func(value, values):
#     var = 1
#     values[0] = 44
# t = 3
# v = [1, 2, 3]
# func(t, v)
# print(t, v[0])""",
                    "explanations": [
                        "Line 1 creates a function with one normal value and one list.",
                        "Line 2 makes a local variable, but it is not used later.",
                        "Line 3 changes the first element of the list to 44.",
                        "Line 4 stores an integer in t.",
                        "Line 5 stores a list in v.",
                        "Line 6 calls the function.",
                        "Line 7 prints t and the first value of v.",
                    ],
                    "output": "t stays 3, but v[0] becomes 44 because the list itself was modified.",
                },
                {
                    "title": "Mutable default list argument",
                    "source": "list.py lines 9-14",
                    "purpose": "This code shows a common Python behavior where one default list is reused across function calls.",
                    "code": """def f(i, values = []):
    values.append(i)
    return values
print(f(1))
print(f(2))
print(f(3))""",
                    "explanations": [
                        "Line 1 defines a function with a default empty list.",
                        "Line 2 adds the new value i into that list.",
                        "Line 3 returns the same list.",
                        "Line 4 calls the function with 1, so the list becomes [1].",
                        "Line 5 calls the function again with 2, but Python reuses the same list, so it becomes [1, 2].",
                        "Line 6 again reuses the same list, so it becomes [1, 2, 3].",
                    ],
                    "note": "This is valid Python, but it can confuse beginners because the list is shared between calls.",
                    "output": "Final output: [1], then [1, 2], then [1, 2, 3].",
                },
            ],
        },
        {
            "title": "3. Linked List",
            "summary": "This chapter explains how nodes are connected and how a linked list grows step by step.",
            "sections": [
                {
                    "title": "Manual node creation and traversal",
                    "source": "linkedlist.py lines 26-49",
                    "purpose": "This example builds a linked list by hand and then prints each node.",
                    "code": """# class Node:
#     def __init__(self, data):
#         self.data = data
#         self.next = None
#
# class LinkedList:
#     def __init__(self):
#         self.head = None
#
# LinkedList = LinkedList()
# LinkedList.head = Node(5)
# second = Node(10)
# third = Node(15)
# fourth = Node(20)
#
# LinkedList.head.next = second
# second.next = third
# third.next = fourth
#
# current = LinkedList.head
# while current is not None:
#     print(current.data)
#     current = current.next""",
                    "explanations": [
                        "The Node class stores one value in data and the next node address in next.",
                        "The LinkedList class starts with head = None, which means the list is empty.",
                        "A linked list object is created, and the first node with value 5 becomes the head.",
                        "More nodes are created with values 10, 15, and 20.",
                        "The next links are connected one by one: 5 -> 10 -> 15 -> 20.",
                        "current starts from the head node.",
                        "The while loop prints the current node value and then moves to the next node.",
                        "The loop stops when current becomes None.",
                    ],
                    "dry_run": "Traversal order is 5, then 10, then 15, then 20.",
                    "output": "This is the basic idea of a singly linked list: every node points to the next node.",
                },
                {
                    "title": "Add nodes at the end and print the list",
                    "source": "linkedlist.py lines 105-148",
                    "purpose": "This cleaner version adds nodes one by one at the end of the list.",
                    "code": """# class Node:
#     def __init__(self, data):
#         self.data = data
#         self.next = None
#
# class LinkedList:
#     def __init__(self):
#         self.head = None
#
#     def add_node(self, data):
#         new_node = Node(data)
#         if self.head is None:
#             self.head = new_node
#         else:
#             temp = self.head
#             while temp.next is not None:
#                 temp = temp.next
#             temp.next = new_node
#
#     def print_list(self):
#         temp = self.head
#         while temp is not None:
#             print(temp.data, end=" -> ")
#             temp = temp.next
#         print("None")""",
                    "explanations": [
                        "The Node class again stores a value and a pointer to the next node.",
                        "The LinkedList class starts with an empty head.",
                        "add_node() creates a new node from the given data.",
                        "If the list is empty, the new node becomes the head.",
                        "Otherwise, temp starts from the head and moves until the last node.",
                        "After reaching the last node, temp.next is linked with the new node.",
                        "print_list() again starts from the head.",
                        "The loop prints each value followed by an arrow.",
                        'At the end it prints "None" to show the list is finished.',
                    ],
                    "output": 'If we add 10, 20, 30, 40, the list prints as "10 -> 20 -> 30 -> 40 -> None".',
                },
                {
                    "title": "Active linked list menu program",
                    "source": "linkedlist.py lines 150-212",
                    "purpose": "This is the active linked list program in the file. It tries to support insert and display operations.",
                    "code": """import sys
class Node:
    def __init__(self, data):
        self.data = data
        self.next = None

class LinkedList:
    def __init__(self):
        self.head = None
        self.tail = None

    def addNode(self, value):
        self.node = Node(value)
        if self.head is None:
            self.head = self.node
            self.tail = self.node
        else:
            self.tail.next = self.node
            self.tail = self.node

    def addNodeBeginning(self, values):
        print("Add node beginning")
        self.node = Node(value)
        if self.head is None:
            self.head = self.node
            self.tail = self.node
        else:
            self.node.next = self.head
            self.head = self.node

    def displayNode(self):
        while self.head is not None:
            print(self.head.data, '|', '->', end='')
            self.head = self.head.next""",
                    "explanations": [
                        "import sys is used later for exiting the menu program.",
                        "The Node class stores one value and the next link.",
                        "The LinkedList class uses both head and tail pointers.",
                        "addNode() creates a new node.",
                        "If the list is empty, both head and tail point to that first node.",
                        "Otherwise, the current tail is linked to the new node, and tail moves forward.",
                        "addNodeBeginning() is meant to insert at the beginning.",
                        "The print line inside addNodeBeginning() only shows a message.",
                        "The if-else logic in addNodeBeginning() is correct in idea: new node should point to the old head, then head should move to the new node.",
                        "displayNode() prints nodes one by one.",
                        "The while loop moves through next links.",
                    ],
                    "note": "This active code has mistakes as written: the parameter name is values but Node(value) uses value, addNodeBetween is unfinished, displayNode changes self.head directly, and the menu calls object.addNodeBeggining(value) with a spelling mismatch.",
                    "output": "Main linked list idea: tail helps fast insert at the end, and head is used to start traversal.",
                },
            ],
        },
        {
            "title": "4. Stack",
            "summary": "This chapter explains the stack code present in stack.py. The stack follows LIFO: last in, first out.",
            "sections": [
                {
                    "title": "Stack using linked list",
                    "source": "stack.py lines 197-259",
                    "purpose": "This code creates a stack with linked list nodes and basic operations like push, pop, and peek.",
                    "code": """# class Node:
#     def __init__(self, value=None):
#         self.value = value
#         self.next = None
#
# class LinkedList:
#     def __init__(self):
#         self.head = None
#
# class Stack:
#     def __init__(self):
#         self.LinkedList = LinkedList()
#
#     def isEmpty(self):
#         if self.LinkedList.head == None:
#             return True
#         else:
#             return False
#
#     def push(self, value):
#         node = Node(value)
#         node.next = self.LinkedList.head
#         self.LinkedList.head = node
#
#     def pop(self):
#         if self.isEmpty():
#             return "There is no element in the stack"
#         else:
#             nodeValue = self.LinkedList.head.value
#             self.LinkedList.head = self.LinkedList.head.next
#             return nodeValue
#
#     def peek(self):
#         if self.isEmpty():
#             return "There is no element in the stack"
#         else:
#             return self.LinkedList.head.value
#
#     def delete(self):
#         self.LinkedList.head = None""",
                    "explanations": [
                        "The Node class stores one stack value and the link to the next node.",
                        "The LinkedList class keeps only head, because the top of the stack is enough here.",
                        "The Stack class stores one linked list object inside it.",
                        "isEmpty() checks whether head is None.",
                        "push() creates a new node.",
                        "The new node points to the old head, so it goes to the top of the stack.",
                        "Then head is moved to the new node.",
                        "pop() first checks whether the stack is empty.",
                        "If it is not empty, it saves the top value, moves head to the next node, and returns the removed value.",
                        "peek() only reads the top value without deleting it.",
                        "delete() clears the whole stack by making head None.",
                    ],
                    "dry_run": "If we push 10, then 20, then 30, the top becomes 30. A pop removes 30 first.",
                    "output": "This is LIFO behavior: the most recent value comes out first.",
                }
            ],
        },
        {
            "title": "5. Queue",
            "summary": "This chapter explains both queue versions found in the workspace: one with a normal list and one with linked list nodes.",
            "sections": [
                {
                    "title": "Queue using Python list and menu",
                    "source": "queue.py lines 1-57",
                    "purpose": "This code creates a queue with a fixed size and a menu to enqueue, dequeue, display, delete, and peek.",
                    "code": """import sys
class Queue:
    def __init__(self, size):
        self.myQueue = []
        self.queueSize = size

    def isFull(self):
        if len(self.myQueue) == self.queueSize:
            return True
        else:
            return False

    def enqueue(self, value):
        if not self.isFull():
            self.myQueue.append(value)
            print("Element enqueued: ", value)
        else:
            print("Queue is full. Cannot enqueue.")""",
                    "explanations": [
                        "import sys is used for exiting the program.",
                        "The Queue class stores items in a Python list named myQueue.",
                        "queueSize stores the maximum number of items allowed.",
                        "isFull() compares the current number of elements with the allowed size.",
                        "If the queue is not full, enqueue() adds the new value at the end of the list.",
                        "If the queue is full, it prints an error message.",
                    ],
                    "note": "In the menu part below this class, the code directly uses queue.myQueue.append() instead of calling queue.enqueue(), but the queue logic still stays understandable.",
                    "output": "Queue rule: insert at the rear, remove from the front.",
                },
                {
                    "title": "Queue menu operations",
                    "source": "queue.py lines 19-57",
                    "purpose": "This part runs the queue interactively.",
                    "code": """size = int(input("Enter the size of the queue: "))
queue = Queue(size)
print("Queue created with size: ", queue.queueSize)
while True:
    print("1. Enqueue")
    print("2. Dequeue")
    print("3. Display")
    print("4. Delete Queue")
    print("5. peek")
    print("6. Exit")""",
                    "explanations": [
                        "The first line asks the user for queue size.",
                        "A queue object is then created using that size.",
                        "The next print confirms the queue size.",
                        "while True means the menu keeps running until the user exits.",
                        "The print lines show the available queue operations.",
                        "Choice 1 adds an element if the queue still has space.",
                        "Choice 2 removes the front element using pop(0).",
                        "Choice 3 prints the full queue list.",
                        "Choice 4 clears the queue.",
                        "Choice 5 shows the front element.",
                        "Choice 6 exits the program using sys.exit().",
                    ],
                    "output": "This program behaves like a simple FIFO queue: first in, first out.",
                },
                {
                    "title": "Queue using linked list nodes",
                    "source": "stack.py lines 261-314",
                    "purpose": "This second queue version stores queue elements in linked list form.",
                    "code": """class Node:
    def __init__(self, value=None):
        self.value = value
        self.next = None

class LinkedList:
    def __init__(self):
        self.head = None

    def __iter__(self):
        current = self.head
        while current:
            yield current.value
            current = current.

class Queue:
    def __init__(self):
        self.LinkedList = LinkedList()

    def enqueue(self, value):
        node = Node(value)
        if self.LinkedList.head is None:
            self.LinkedList.head = node
        else:
            current = self.LinkedList.head
            while current.next:
                current = current.next
            current.next = node""",
                    "explanations": [
                        "The Node class stores one queue value and the next pointer.",
                        "The LinkedList class stores the front node in head.",
                        "__iter__() is meant to help loop through the linked list values.",
                        "Queue stores one LinkedList object inside it.",
                        "enqueue() creates a new node for the incoming value.",
                        "If the queue is empty, the first node becomes the head.",
                        "Otherwise, current starts from the head and moves to the last node.",
                        "The last node is then linked to the new node.",
                        "In the remaining methods of the file, dequeue() removes the front node and peek() reads the front value.",
                    ],
                    "note": "This code has a syntax error in __iter__(): the line current = current. is incomplete. The queue methods below it are logically fine, but the file cannot run until that line is fixed.",
                    "output": "This version also follows FIFO, but stores items as linked nodes instead of a Python list.",
                },
            ],
        },
        {
            "title": "6. Tree",
            "summary": "This chapter explains the tree examples from tree.py and ignores the unrelated CSV code at the bottom.",
            "sections": [
                {
                    "title": "General tree with children list",
                    "source": "tree.py lines 1-31",
                    "purpose": "This code builds a normal tree where one node can have many children.",
                    "code": """# class Tree:
#     def __init__(self, data):
#         self.data = data
#         self.child = []
#
#     def __str__(self, level=0):
#         ret = "\\t" * level + repr(self.data) + "\\n"
#         for child in self.child:
#             ret += child.__str__(level + 1)
#         return ret
#
#     def addChild(self, object):
#         self.child.append(object)
#
# rootNode = Tree("Drinks")
# hot = Tree("Hot")
# cold = Tree("Cold")
# Tea = Tree("Tea")
# Coffee = Tree("Coffee")
# rootNode.addChild(hot)
# rootNode.addChild(cold)
# hot.addChild(Tea)
# hot.addChild(Coffee)""",
                    "explanations": [
                        "The Tree class stores one data value and a list named child.",
                        "child = [] means one node can have many child nodes.",
                        "__str__() is used to print the tree in a structured way.",
                        "It adds tabs based on level so children appear under parents.",
                        "addChild() simply appends a child node to the child list.",
                        'rootNode = Tree("Drinks") creates the main root node.',
                        "Other lines create child nodes like Hot, Cold, Tea, and Coffee.",
                        "The addChild calls connect the nodes into a tree shape.",
                    ],
                    "output": 'This is a multi-child tree: "Drinks" is the root, and "Hot" and "Cold" are children.',
                },
                {
                    "title": "Binary tree and preorder traversal",
                    "source": "tree.py lines 35-73",
                    "purpose": "This code builds a binary tree and prints it using preorder traversal.",
                    "code": """# class Node:
#     def __init__(self, value):
#         self.value = value
#         self.left = None
#         self.right = None
#
# N1 = Node("N1")
# N2 = Node("N2")
# N3 = Node("N3")
# N4 = Node("N4")
#
# N1.left = N2
# N1.right = N3
# N2.left = N4
#
# def preorder(node):
#     if node:
#         print(node.value, end=" ")
#         preorder(node.left)
#         preorder(node.right)""",
                    "explanations": [
                        "The Node class stores one value and two links: left and right.",
                        "left = None and right = None mean no child is connected yet.",
                        "N1, N2, N3, and N4 create tree nodes.",
                        "N1.left = N2 connects N2 to the left of N1.",
                        "N1.right = N3 connects N3 to the right of N1.",
                        "N2.left = N4 connects N4 to the left of N2.",
                        "preorder(node) is a recursive traversal function.",
                        "It first prints the current node, then visits the left subtree, then the right subtree.",
                    ],
                    "dry_run": "For this small tree, preorder order starts from root first, then left side, then right side.",
                    "output": "Preorder rule is Root -> Left -> Right.",
                },
                {
                    "title": "Binary Search Tree operations",
                    "source": "tree.py lines 122-227",
                    "purpose": "This code shows how a BST stores smaller values on the left and bigger values on the right.",
                    "code": """# class BSTNode:
#     def __init__(self, data):
#         self.data = data
#         self.leftChild = None
#         self.rightChild = None
#
# def insertNode(rootNode, nodeValue):
#     if rootNode.data is None:
#         rootNode.data = nodeValue
#     elif nodeValue <= rootNode.data:
#         if rootNode.leftChild is None:
#             rootNode.leftChild = BSTNode(nodeValue)
#         else:
#             insertNode(rootNode.leftChild, nodeValue)
#     else:
#         if rootNode.rightChild is None:
#             rootNode.rightChild = BSTNode(nodeValue)
#         else:
#             insertNode(rootNode.rightChild, nodeValue)
#
# def searchNode(rootNode, nodeValue):
#     if rootNode is None:
#         print("The value is not found")
#     elif rootNode.data == nodeValue:
#         print("The value is found")""",
                    "explanations": [
                        "BSTNode stores one data value and two child links.",
                        "insertNode() adds a value into the correct place in the BST.",
                        "If the new value is smaller or equal, it goes to the left side.",
                        "If the new value is bigger, it goes to the right side.",
                        "If the needed child position is empty, a new BSTNode is created there.",
                        "Otherwise, insertNode() calls itself again on the next child.",
                        "searchNode() checks whether the required value exists.",
                        "If the current node is None, the value is not found.",
                        "If the current node data matches, the value is found.",
                        "The traversal functions later in the file print the tree in preorder, inorder, and postorder style.",
                    ],
                    "dry_run": "If root is 70 and we insert 50, it goes left. If we insert 90, it goes right.",
                    "output": "BST makes searching easier because left side has smaller values and right side has bigger values.",
                },
            ],
        },
        {
            "title": "7. Graph",
            "summary": "This chapter explains the graph class in graph.py and ignores the unrelated student and inheritance code below it.",
            "sections": [
                {
                    "title": "Graph using adjacency list",
                    "source": "graph.py lines 1-57",
                    "purpose": "This code builds an undirected graph using a dictionary of lists.",
                    "code": """class Graph:
    def __init__(self):
        self.adjacency_list = {}

    def add_vertex(self, vertex):
        if vertex not in self.adjacency_list.keys():
            self.adjacency_list[vertex] = []
            return True
        return False

    def add_edge(self, vertex1, vertex2):
        if vertex1 in self.adjacency_list.keys() and vertex2 in self.adjacency_list.keys():
            self.adjacency_list[vertex1].append(vertex2)
            self.adjacency_list[vertex2].append(vertex1)
            return True
        return False

    def remove_vertex(self, vertex):
        if vertex in self.adjacency_list.keys():
            for other_vertex in self.adjacency_list[vertex]:
                self.adjacency_list[other_vertex].remove(vertex)
            del self.adjacency_list[vertex]
            return True
        return False""",
                    "explanations": [
                        "The Graph class stores data in adjacency_list, which is a dictionary.",
                        "Each key is a vertex name, and each value is a list of connected vertices.",
                        "add_vertex() first checks whether the vertex already exists.",
                        "If not, it creates an empty list for that vertex.",
                        "add_edge() checks that both vertices already exist.",
                        "Because this is an undirected graph, each vertex is added to the other vertex list.",
                        "remove_vertex() first checks whether the vertex exists.",
                        "It then removes that vertex name from every connected neighbor.",
                        "After that, it deletes the vertex itself from the dictionary.",
                    ],
                    "dry_run": "If A is connected to B and C, then A's list stores [B, C]. Because the graph is undirected, B also stores A and C also stores A when connected.",
                    "output": "Adjacency list is a simple way to store graph connections.",
                },
                {
                    "title": "Graph object usage",
                    "source": "graph.py lines 36-57",
                    "purpose": "This part creates vertices, connects them, prints the graph, removes one vertex, and prints again.",
                    "code": """my_graph = Graph()
my_graph.add_vertex("A")
my_graph.add_vertex("B")
my_graph.add_vertex("C")
my_graph.add_vertex("D")
my_graph.add_vertex("E")

my_graph.add_edge("A", "B")
my_graph.add_edge("A", "C")
my_graph.add_edge("A", "D")
my_graph.add_edge("B", "E")
my_graph.add_edge("C", "D")
my_graph.add_edge("D", "E")

print("Before removing vertex:")
my_graph.print_graph()
my_graph.remove_vertex("D")
print("\\nAfter removing vertex D:")
my_graph.print_graph()""",
                    "explanations": [
                        "The first line creates a graph object.",
                        'The next five lines add vertices A, B, C, D, and E.',
                        "The add_edge lines connect the vertices in pairs.",
                        "print_graph() shows the adjacency list before deletion.",
                        'remove_vertex("D") deletes D and also removes D from other neighbor lists.',
                        "The final print_graph() shows the graph after D is removed.",
                    ],
                    "output": "This demonstrates create, connect, display, and delete operations in a graph.",
                },
            ],
        },
    ]


if __name__ == "__main__":
    build_document()
